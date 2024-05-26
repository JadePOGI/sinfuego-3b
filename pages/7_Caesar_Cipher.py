import streamlit as st
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io
import docx

st.header("Welcome to Caesar Cipher!🔐")
st.write("Caesar Cipher Technique is the simple and easy method of encryption technique.")

def caesar_cipher(text, shift):
    encrypted_text = ""
    for char in text:
        if char.isalpha():  # Check if the character is a letter
            shifted = ord(char) + shift
            if char.islower():
                if shifted > ord('z'):
                    shifted -= 26
                elif shifted < ord('a'):
                    shifted += 26
            elif char.isupper():
                if shifted > ord('Z'):
                    shifted -= 26
                elif shifted < ord('A'):
                    shifted += 26
            encrypted_text += chr(shifted)
        else:
            encrypted_text += char
    return encrypted_text

def caesar_decipher(text, shift):
    return caesar_cipher(text, -shift)

def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
    return text

def read_docx(file):
    doc = docx.Document(io.BytesIO(file.read()))
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)

def create_pdf_from_text(text):
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    text_object = c.beginText(40, 750)
    text_object.setFont("Helvetica", 12)

    lines = text.split('\n')
    for line in lines:
        text_object.textLine(line)
    
    c.drawText(text_object)
    c.showPage()
    c.save()
    pdf_buffer.seek(0)
    return pdf_buffer

def create_docx_from_text(text):
    docx_buffer = io.BytesIO()
    doc = docx.Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

genre = st.radio("Choose Input:", ["Text", "File"])

if genre == 'Text':
    plaintext = st.text_area('Enter your text to Encrypt or Decrypt')
    shift = 5

    action = st.radio('Choose one', options=('Encrypt', 'Decrypt'))
    if st.button('Submit'):
        if action == 'Encrypt':
            encrypted_text = caesar_cipher(plaintext, shift)
            st.write("Encrypted:", encrypted_text)
        elif action == 'Decrypt':
            decrypted_text = caesar_decipher(plaintext, shift)
            st.write("Decrypted:", decrypted_text)

elif genre == 'File':
    action = st.radio('Choose one', options=('Encrypt', 'Decrypt'))
    if action == 'Encrypt':
        uploaded_file = st.file_uploader("Choose a file to Encrypt", type=["pdf", "txt", "docx"])
        if uploaded_file is not None:
            file_contents = None
            file_type = uploaded_file.type.split('/')[-1]
            original_file_name = uploaded_file.name.replace(".encrypted", "")  # Remove .encrypted extension if present

            if file_type == "pdf":
                file_contents = read_pdf(uploaded_file)
            elif file_type == "plain":
                file_contents = uploaded_file.getvalue().decode("utf-8")
            elif file_type == "vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_contents = read_docx(uploaded_file)
            
            if file_contents is None:
                st.write("Failed to extract file contents. Please upload a valid file type.")
            else:
                st.write("File contents extracted successfully.")
                shift = 5
                encrypted_text = caesar_cipher(file_contents, shift)
                if file_type == "pdf":
                    encrypted_pdf = create_pdf_from_text(encrypted_text)
                    st.download_button(
                        label="Download Encrypted File",
                        data=encrypted_pdf.getvalue(),
                        file_name=original_file_name + ".encrypted",
                        mime="application/pdf"
                    )
                elif file_type == "plain":
                    st.download_button(
                        label="Download Encrypted File",
                        data=encrypted_text,
                        file_name=original_file_name + ".encrypted",
                        mime="text/plain"
                    )
                elif file_type == "vnd.openxmlformats-officedocument.wordprocessingml.document":
                    encrypted_docx = create_docx_from_text(encrypted_text)
                    st.download_button(
                        label="Download Encrypted File",
                        data=encrypted_docx.getvalue(),
                        file_name=original_file_name + ".encrypted",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    elif action == 'Decrypt':
        uploaded_encrypted_file = st.file_uploader("Choose an encrypted file to Decrypt", type=["pdf", "txt", "docx"])
        if uploaded_encrypted_file is not None:
            file_contents = None
            file_type = uploaded_encrypted_file.type.split('/')[-1]
            original_file_name = uploaded_encrypted_file.name.replace(".encrypted", "")  # Remove .encrypted extension if present

            if file_type == "pdf":
                file_contents = read_pdf(uploaded_encrypted_file)
            elif file_type == "plain":
                file_contents = uploaded_encrypted_file.getvalue().decode("utf-8")
            elif file_type == "vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_contents = read_docx(uploaded_encrypted_file)
